"""Fills the real proforma-invoice .docx template used to buy from Chinese
suppliers, then converts it to PDF via LibreOffice headless.

The buyer/seller/bank blocks and the single line-item's product name are
baked into a per-company template under assets/invoice_templates/ — only
Tarih, Contract No, Unit Price and Total Price change per invoice.
Quantity is never entered — it's derived as Total Price / Unit Price, and
written into both the line-item row and the totals row. Content is never
translated (the template itself is already bilingual ZH/EN).

Adding a new seller: drop its filled-in-once .docx into
invoice_templates/ and add an entry to INVOICE_TEMPLATES with the exact
fixed prefixes/table coordinates found in that template.
"""

from __future__ import annotations

import io
import os
from dataclasses import dataclass

import docx

from .docx_pdf import (
    docx_bytes_to_pdf,
    ensure_cjk_font_installed,
    force_east_asian_font,
    remove_paragraph,
    replace_after_prefix,
    set_paragraph_text,
)

ASSETS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "invoice_templates")

CJK_FALLBACK_FONT = "Noto Sans SC"

CURRENCY_SYMBOLS = {"USD": "$", "EUR": "€", "CNY": "¥", "TRY": "₺"}


@dataclass
class InvoiceTemplate:
    key: str
    label: str
    file: str
    date_prefix: str
    contract_prefix: str
    contract_continuation_paragraph_index: int
    line_item_row: int
    quantity_col: int
    unit_price_col: int
    total_price_col: int
    totals_row: int
    unit: str = "Kg"


INVOICE_TEMPLATES: dict[str, InvoiceTemplate] = {
    "urumqi_yilu_qixin": InvoiceTemplate(
        key="urumqi_yilu_qixin",
        label="Urumqi Yilu Qixin Trading Co., Ltd",
        file="urumqi_yilu_qixin.docx",
        date_prefix="Date:",
        contract_prefix="Contract:",
        contract_continuation_paragraph_index=3,
        line_item_row=2,
        quantity_col=2,
        unit_price_col=3,
        total_price_col=4,
        totals_row=4,
    ),
}


def list_companies() -> list[dict]:
    return [{"key": t.key, "label": t.label} for t in INVOICE_TEMPLATES.values()]


def _format_quantity(quantity: float) -> str:
    if abs(quantity - round(quantity)) < 0.005:
        return f"{round(quantity):,}"
    return f"{quantity:,.2f}"


def _cell_text_paragraph(cell):
    """Returns the paragraph actually holding the value text — these cells
    have a leading empty paragraph before the one with content."""
    for p in cell.paragraphs:
        if p.text.strip():
            return p
    return cell.paragraphs[-1]


def fill_template_docx(
    template: InvoiceTemplate,
    tarih: str,
    contract_no: str,
    unit_price: float,
    total_price: float,
    currency: str = "USD",
) -> bytes:
    path = os.path.join(ASSETS_DIR, template.file)
    doc = docx.Document(path)

    symbol = CURRENCY_SYMBOLS.get(currency, currency)
    quantity = total_price / unit_price if unit_price else 0
    quantity_text = f"{_format_quantity(quantity)} {template.unit}"
    unit_price_str = f"{unit_price:.1f}" if unit_price == int(unit_price) else f"{unit_price:.2f}"
    unit_price_text = f"{symbol} {unit_price_str}"
    total_price_text = f"{symbol} {round(total_price)}"

    paragraphs = doc.paragraphs
    for p in paragraphs:
        replace_after_prefix(p, template.date_prefix, tarih)
        replace_after_prefix(p, template.contract_prefix, contract_no)

    continuation = paragraphs[template.contract_continuation_paragraph_index]
    remove_paragraph(continuation)

    table = doc.tables[0]
    line_row = table.rows[template.line_item_row]
    set_paragraph_text(_cell_text_paragraph(line_row.cells[template.quantity_col]), quantity_text)
    set_paragraph_text(_cell_text_paragraph(line_row.cells[template.unit_price_col]), unit_price_text)
    set_paragraph_text(_cell_text_paragraph(line_row.cells[template.total_price_col]), total_price_text)

    totals_row = table.rows[template.totals_row]
    set_paragraph_text(_cell_text_paragraph(totals_row.cells[template.quantity_col]), quantity_text)
    set_paragraph_text(_cell_text_paragraph(totals_row.cells[template.total_price_col]), total_price_text)

    force_east_asian_font(doc, CJK_FALLBACK_FONT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_invoice_pdf(
    company_key: str,
    tarih: str,
    contract_no: str,
    unit_price: float,
    total_price: float,
    currency: str = "USD",
) -> bytes:
    ensure_cjk_font_installed()
    template = INVOICE_TEMPLATES[company_key]
    docx_bytes = fill_template_docx(template, tarih, contract_no, unit_price, total_price, currency)
    return docx_bytes_to_pdf(docx_bytes)
