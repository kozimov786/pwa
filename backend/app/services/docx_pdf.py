"""Shared docx->PDF conversion via LibreOffice headless, and small
run-level text-editing helpers for filling docx templates while
preserving their original layout/formatting."""

from __future__ import annotations

import os
import shutil
import subprocess
import tempfile

from docx.oxml.ns import qn

_FONTS_ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "assets", "fonts")
_LIBREOFFICE_FONT_DIR = "/Applications/LibreOffice.app/Contents/Resources/fonts/truetype"
_cjk_font_checked = False


def ensure_cjk_font_installed():
    """This LibreOffice cask ships with no fontconfig.conf, so it can only
    see the small bundled font set under its own Resources/fonts/truetype —
    it never finds any macOS system font, CJK or otherwise, and silently
    drops unrenderable glyphs instead of falling back. Copy our bundled
    Noto Sans SC into that directory (once) so CJK text in filled templates
    actually renders. On Linux this isn't needed if a system CJK package
    (e.g. fonts-noto-cjk) is installed instead."""
    global _cjk_font_checked
    if _cjk_font_checked:
        return
    _cjk_font_checked = True
    src = os.path.join(_FONTS_ASSETS_DIR, "NotoSansSC.ttf")
    dst = os.path.join(_LIBREOFFICE_FONT_DIR, "NotoSansSC.ttf")
    if os.path.isdir(_LIBREOFFICE_FONT_DIR) and os.path.exists(src) and not os.path.exists(dst):
        shutil.copy(src, dst)


def _has_cjk(text: str) -> bool:
    return any("一" <= ch <= "鿿" for ch in text)


def force_east_asian_font(doc, font_name: str):
    """LibreOffice's headless PDF export does not substitute a fallback
    font for missing glyphs the way Word does — if a run's eastAsia font
    (often left at the document default, e.g. 'Times New Roman', which has
    no CJK glyphs) can't render its characters, they come out blank rather
    than as tofu boxes. Force just the runs that actually contain CJK text
    to a font guaranteed to have real CJK glyphs on this machine, leaving
    the document's Latin typography untouched."""

    def apply(run):
        if not _has_cjk(run.text):
            return
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), font_name)

    for p in doc.paragraphs:
        for r in p.runs:
            apply(r)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        apply(r)


def set_paragraph_text(paragraph, new_text: str):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_after_prefix(paragraph, prefix: str, new_value: str) -> bool:
    if not paragraph.text.startswith(prefix):
        return False
    set_paragraph_text(paragraph, prefix + new_value)
    return True


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def docx_bytes_to_pdf(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "form.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        subprocess.run(
            [
                "soffice", "--headless", "--norestore",
                "--convert-to", "pdf", "--outdir", tmp, docx_path,
            ],
            check=True,
            capture_output=True,
            timeout=60,
        )

        pdf_path = os.path.join(tmp, "form.pdf")
        with open(pdf_path, "rb") as f:
            return f.read()
