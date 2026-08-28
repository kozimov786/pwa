"""Fills the real Vakıfbank wire-transfer .docx form used to pay Chinese
suppliers, then converts the filled form to PDF via LibreOffice headless.

The form's fixed fields (beneficiary name/address/bank/IBAN/SWIFT) are
baked into a per-company template under assets/bank_templates/ — only the
transaction-specific fields change on every payment: Tarih (date), Valör
(value date), the transferred amount+currency (with its Turkish
amount-in-words line), and the Swift Açıklaması / Contract No. Content is
never translated — this document goes to a Turkish bank and must stay in
Turkish regardless of the app's UI language.

Adding a new beneficiary company: drop its filled-in-once .docx into
bank_templates/ and add an entry to COMPANY_TEMPLATES with the exact fixed
prefix text found in that template for each of the four dynamic fields.
"""

from __future__ import annotations

import io
import os
import subprocess
import tempfile
from dataclasses import dataclass

import docx

from .turkish_numbers import number_to_words_tr, turkish_upper

ASSETS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "bank_templates")

CURRENCY_SYMBOLS = {"USD": "$", "EUR": "€", "CNY": "¥", "TRY": "₺"}
CURRENCY_WORDS_TR_CAPS = {"USD": "DOLAR", "EUR": "AVRO", "CNY": "YUAN", "TRY": "TL"}


@dataclass
class CompanyTemplate:
    key: str
    label: str
    file: str
    date_prefix: str
    amount_paragraph_is_standalone: bool
    words_prefix: str
    valor_prefix: str
    contract_prefix: str


COMPANY_TEMPLATES: dict[str, CompanyTemplate] = {
    "urumqi_yilu_qixin": CompanyTemplate(
        key="urumqi_yilu_qixin",
        label="Urumqi Yilu Qixin Trading Co., Ltd",
        file="urumqi_yilu_qixin.docx",
        date_prefix="TÜRKİYE VAKIFLAR BANKASI T.A.O\tTarih: ",
        amount_paragraph_is_standalone=True,
        words_prefix="(Yazı ve Rakamla)\t: ",
        valor_prefix="Valör\t: ",
        contract_prefix="Swift Açıklaması\t: Contract No : ",
    ),
}


def list_companies() -> list[dict]:
    return [{"key": c.key, "label": c.label} for c in COMPANY_TEMPLATES.values()]


def amount_to_words_caps(amount: float, currency: str) -> str:
    major = int(amount)
    minor = round((amount - major) * 100)
    currency_word = CURRENCY_WORDS_TR_CAPS.get(currency, currency)
    text = f"{turkish_upper(number_to_words_tr(major))} {currency_word}"
    if minor:
        text += f" {turkish_upper(number_to_words_tr(minor))} KURUŞ"
    return text


def _set_paragraph_text(paragraph, new_text: str):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_after_prefix(paragraph, prefix: str, new_value: str) -> bool:
    if not paragraph.text.startswith(prefix):
        return False
    _set_paragraph_text(paragraph, prefix + new_value)
    return True


def fill_template_docx(
    template: CompanyTemplate,
    tarih: str,
    valor_tarihi: str,
    amount: float,
    currency: str,
    contract_no: str,
) -> bytes:
    path = os.path.join(ASSETS_DIR, template.file)
    doc = docx.Document(path)

    symbol = CURRENCY_SYMBOLS.get(currency, currency)
    amount_line = f"{int(amount)} {symbol}"
    amount_words = amount_to_words_caps(amount, currency)

    for p in doc.paragraphs:
        _replace_after_prefix(p, template.date_prefix, tarih)
        _replace_after_prefix(p, template.words_prefix, amount_words)
        _replace_after_prefix(p, template.valor_prefix, valor_tarihi)
        _replace_after_prefix(p, template.contract_prefix, contract_no)
        if template.amount_paragraph_is_standalone and p.text.strip().endswith(("$", "€", "¥", "₺")):
            _set_paragraph_text(p, amount_line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_bytes_to_pdf(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "form.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        subprocess.run(
            [
                "soffice", "--headless", "--norestore",
                "--convert-to", "pdf", "--outdir", tmp, docx_path,
            ],
            check=True,
            capture_output=True,
            timeout=60,
        )

        pdf_path = os.path.join(tmp, "form.pdf")
        with open(pdf_path, "rb") as f:
            return f.read()


def generate_transfer_pdf(
    company_key: str,
    tarih: str,
    valor_tarihi: str,
    amount: float,
    currency: str,
    contract_no: str,
) -> bytes:
    template = COMPANY_TEMPLATES[company_key]
    docx_bytes = fill_template_docx(template, tarih, valor_tarihi, amount, currency, contract_no)
    return docx_bytes_to_pdf(docx_bytes)
